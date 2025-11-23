import hashlib
import logging
from pathlib import Path
from typing import List

from langchain.docstore.document import Document

# 配置日志
logger = logging.getLogger(__name__)

def load_documents(self) -> List[Document]:
    """
    加载文档数据，支持RagFlow处理的多种文件格式

    Returns:
        加载的文档列表
    """
    logger.info(f"正在从 {self.data_path} 加载文档...")

    # 支持的文件扩展名
    supported_extensions = {
        # 文本类文件
        '.md', '.txt', '.rtf',
        # 文档类文件
        '.pdf', '.doc', '.docx',
        # 表格类文件
        '.xls', '.xlsx', '.csv',
        # 演示文稿类文件
        '.ppt', '.pptx',
        # 网页文件
        '.html', '.htm',
        # 邮件文件
        '.eml', '.msg',
        # 图片文件（支持OCR）
        '.png', '.jpg', '.jpeg', '.tiff', '.bmp'
    }

    documents = []
    data_path_obj = Path(self.data_path)

    # 使用RagFlow的文档加载器
    try:
        from ragflow import RAGFlow, DocumentIngest

        # 初始化RagFlow（根据你的配置调整）
        ragflow = RAGFlow()

        # 获取所有支持的文件
        all_files = []
        for ext in supported_extensions:
            all_files.extend(data_path_obj.rglob(f"*{ext}"))

        logger.info(f"发现 {len(all_files)} 个支持的文件")

        for file_path in all_files:
            try:
                logger.info(f"正在处理文件: {file_path}")

                # 使用RagFlow处理文档
                ingest = DocumentIngest()

                # 根据文件类型选择处理方式
                if file_path.suffix.lower() in ['.pdf', '.doc', '.docx', '.ppt', '.pptx']:
                    # 文档类文件 - 使用RagFlow的高级文档解析
                    result = ingest.parse_document(str(file_path))
                    content = result.text_content
                    metadata = result.metadata if hasattr(result, 'metadata') else {}
                elif file_path.suffix.lower() in ['.xls', '.xlsx', '.csv']:
                    # 表格类文件 - 提取表格数据
                    result = ingest.parse_table(str(file_path))
                    content = result.text_content
                    metadata = result.metadata if hasattr(result, 'metadata') else {}
                elif file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                    # 图片文件 - 使用OCR
                    result = ingest.parse_image(str(file_path), use_ocr=True)
                    content = result.text_content
                    metadata = result.metadata if hasattr(result, 'metadata') else {}
                else:
                    # 文本类文件 - 直接读取
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    metadata = {}

                # 为每个父文档分配确定性的唯一ID
                try:
                    data_root = Path(self.data_path).resolve()
                    relative_path = Path(file_path).resolve().relative_to(data_root).as_posix()
                except Exception:
                    relative_path = Path(file_path).as_posix()
                parent_id = hashlib.md5(relative_path.encode("utf-8")).hexdigest()

                # 创建Document对象
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": str(file_path),
                        "parent_id": parent_id,
                        "doc_type": "parent",
                        "file_extension": file_path.suffix.lower(),
                        "file_size": file_path.stat().st_size,
                        **metadata  # 合并RagFlow提供的元数据
                    }
                )
                documents.append(doc)
                logger.info(f"成功处理文件: {file_path}")

            except Exception as e:
                logger.warning(f"使用RagFlow处理文件 {file_path} 失败: {e}")
                # 回退到基础处理
                try:
                    _fallback_load_file(file_path, documents, self.data_path)
                except Exception as fallback_error:
                    logger.error(f"回退处理也失败 {file_path}: {fallback_error}")

    except ImportError:
        logger.warning("未安装RagFlow，使用基础文档加载模式")
        # 回退到基础模式，只支持常见文本格式
        _load_documents_basic(documents, data_path_obj, self.data_path)

    # 增强文档元数据
    for doc in documents:
        self._enhance_metadata(doc)

    self.documents = documents
    logger.info(f"成功加载 {len(documents)} 个文档")
    return documents

def _fallback_load_file(file_path: Path, documents: List[Document], data_path: str):
    """
    RagFlow处理失败时的回退文件加载方法

    Args:
        file_path: 文件路径
        documents: 文档列表
        data_path: 数据根目录路径
    """
    try:
        # 根据文件类型使用不同的回退策略
        if file_path.suffix.lower() == '.pdf':
            # 使用PyPDF2处理PDF
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        elif file_path.suffix.lower() in ['.docx']:
            # 使用python-docx处理Word文档
            import docx
            doc = docx.Document(str(file_path))
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_path.suffix.lower() in ['.csv']:
            # 使用pandas处理CSV
            import pandas as pd
            df = pd.read_csv(str(file_path))
            content = df.to_string()
        elif file_path.suffix.lower() in ['.txt', '.md', '.rtf', '.html', '.htm']:
            # 文本文件直接读取
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            logger.warning(f"不支持的文件类型进行回退处理: {file_path.suffix}")
            return

        # 为父文档分配ID
        try:
            data_root = Path(data_path).resolve()
            relative_path = Path(file_path).resolve().relative_to(data_root).as_posix()
        except Exception:
            relative_path = Path(file_path).as_posix()
        parent_id = hashlib.md5(relative_path.encode("utf-8")).hexdigest()

        # 创建Document对象
        doc = Document(
            page_content=content,
            metadata={
                "source": str(file_path),
                "parent_id": parent_id,
                "doc_type": "parent",
                "file_extension": file_path.suffix.lower(),
                "file_size": file_path.stat().st_size,
                "loading_method": "fallback"
            }
        )
        documents.append(doc)
        logger.info(f"回退处理成功: {file_path}")

    except ImportError as e:
        logger.warning(f"缺少处理 {file_path.suffix} 文件的依赖库: {e}")
    except Exception as e:
        logger.error(f"回退处理文件 {file_path} 失败: {e}")

def _load_documents_basic(documents: List[Document], data_path_obj: Path, data_path: str):
    """
    基础文档加载模式（当RagFlow不可用时）

    Args:
        documents: 文档列表
        data_path_obj: 数据路径对象
        data_path: 数据路径
    """
    # 基础支持的文件格式
    basic_extensions = {'.md', '.txt', '.html', '.htm'}

    for file_path in data_path_obj.rglob("*"):
        if file_path.suffix.lower() not in basic_extensions:
            continue

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 为每个父文档分配确定性的唯一ID
            try:
                data_root = Path(data_path).resolve()
                relative_path = Path(file_path).resolve().relative_to(data_root).as_posix()
            except Exception:
                relative_path = Path(file_path).as_posix()
            parent_id = hashlib.md5(relative_path.encode("utf-8")).hexdigest()

            # 创建Document对象
            doc = Document(
                page_content=content,
                metadata={
                    "source": str(file_path),
                    "parent_id": parent_id,
                    "doc_type": "parent",
                    "file_extension": file_path.suffix.lower(),
                    "loading_method": "basic"
                }
            )
            documents.append(doc)

        except Exception as e:
            logger.warning(f"基础模式读取文件 {file_path} 失败: {e}")